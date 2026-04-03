from pathlib import Path
from urllib.parse import urlencode

import qrcode
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.lib import colors
from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from config import (
    DOCX_DIR,
    PDF_DIR,
    QR_DIR,
    BASE_VERIFY_URL,
    ORG_NAME,
    ORG_WEBSITE,
    ORG_EMAIL,
    ORG_PHONE,
    ORG_CITY,
    PRESIDENT_NAME,
    PRESIDENT_TITLE,
    REGISTRE_ASSOCIATION,
    COORDINATOR_TITLE,
    LOGO_PATH,
    PRESIDENT_SIGNATURE_PATH,
    STAMP_PATH,
)
from security import generate_signature


AMALIA_REFERENCE = "Référence AMALIA : A2019STR000236 – Volume 97 – Folio 229"
ONG_SIEGE = "Siège : 24 Rue de la Niederbourg, 67400 Illkirch-Graffenstaden, France"
ONG_REPRESENTATION = "Représentée par Monsieur Romuald HOUNYEME, Président"

PRIMARY_COLOR = RGBColor(11, 79, 156)
ACCENT_COLOR = RGBColor(216, 166, 42)
TEXT_COLOR = RGBColor(34, 34, 34)

PDF_PRIMARY = colors.HexColor("#0B4F9C")
PDF_ACCENT = colors.HexColor("#D8A62A")
PDF_TEXT = colors.HexColor("#222222")


def sanitize_reference(ref: str) -> str:
    return ref.replace("/", "_")


def verify_url(reference: str, uid: str, sig: str) -> str:
    query = urlencode(
        {
            "ref": reference,
            "uid": uid,
            "sig": sig,
        }
    )
    return f"{BASE_VERIFY_URL}?{query}"


def generate_qr(reference: str, uid: str, sig: str) -> Path:
    QR_DIR.mkdir(parents=True, exist_ok=True)
    path = QR_DIR / f"{sanitize_reference(reference)}_qr.png"
    qr_url = verify_url(reference, uid, sig)
    img = qrcode.make(qr_url)
    img.save(path)
    return path


def build_text(payload: dict) -> dict:
    civilite = payload.get("civilite", "").strip()
    civilite_nom = f"{civilite} {payload['prenom']} {payload['nom']}".strip()

    return {
        "org_name": ORG_NAME,
        "amalia_reference": AMALIA_REFERENCE,
        "siege": ONG_SIEGE,
        "representation": ONG_REPRESENTATION,
        "president": PRESIDENT_NAME,
        "president_title": PRESIDENT_TITLE,
        "civilite_nom": civilite_nom,
        "adresse": payload.get("adresse") or "Non renseignée",
        "telephone": payload.get("telephone") or "Non renseigné",
        "fonction": payload["fonction"],
        "zone": payload["zone_intervention"],
        "pays": payload.get("pays") or "Bénin",
        "date_emission": payload.get("date_emission", ""),
        "date_expiration": payload.get("date_expiration", ""),
        "ville_signature": payload.get("ville_signature", ORG_CITY),
    }


def mandate_intro_lines(payload: dict) -> list[str]:
    fmt = build_text(payload)
    return [
        f"L’ONG Renaître de Nouveau, organisation régulièrement constituée et inscrite au Registre des Associations sous la référence {fmt['amalia_reference']},",
        f"dont le siège est situé {fmt['siege']},",
        f"{fmt['representation']},",
        "",
        "Certifie et atteste que :",
        "",
        f"{fmt['civilite_nom']},",
        f"demeurant à {fmt['adresse']},",
        f"contact : {fmt['telephone']},",
        "membre de ladite organisation,",
        "",
        "est dûment habilité à représenter l’ONG Renaître de Nouveau en qualité de :",
        "",
        f"{fmt['fonction']}",
        "",
        f"dans la zone d’intervention suivante : {fmt['zone']}, République du {fmt['pays']}.",
    ]


ARTICLE_1_TITLE = "1. Portée du mandat"
ARTICLE_1_TEXT = (
    "Le présent mandat constitue une habilitation officielle conférée par l’ONG, "
    "autorisant son titulaire à intervenir en son nom dans le cadre strict de ses missions, "
    "conformément aux statuts de l’organisation, aux orientations stratégiques validées "
    "et aux instructions des instances dirigeantes."
)

ARTICLE_2_TITLE = "2. Attributions"
ARTICLE_2_BULLETS = [
    "représenter l’ONG auprès des autorités administratives, institutions et partenaires ;",
    "contribuer à la mise en œuvre des activités sociales, humanitaires et communautaires ;",
    "assurer la coordination locale des actions validées par l’organisation ;",
    "rendre compte des activités menées conformément aux procédures internes.",
]

ARTICLE_3_TITLE = "3. Limites de l’habilitation"
ARTICLE_3_TEXT = (
    "Le présent mandat n’emporte aucun pouvoir d’engagement juridique, financier ou contractuel autonome "
    "au nom de l’ONG, sauf autorisation expresse et préalable. "
    "Le titulaire est tenu d’agir dans le strict respect des lois et règlements en vigueur dans le pays "
    "d’intervention, ainsi que des statuts et directives de l’ONG."
)

ARTICLE_4_TITLE = "4. Nature de la mission"
ARTICLE_4_TEXT = (
    "La mission est exercée à titre bénévole. Toute prise en charge de frais éventuels est soumise "
    "aux procédures internes de l’ONG."
)

ARTICLE_5_TITLE = "5. Durée et révocation"
ARTICLE_5_TEXT_TEMPLATE = (
    "Le présent mandat est valable du {date_emission} au {date_expiration}, sauf modification, "
    "suspension ou retrait à tout moment par l’ONG, notamment en cas de manquement."
)

ARTICLE_6_TITLE = "6. Authenticité"
ARTICLE_6_TEXT = (
    "Le présent document est sécurisé par un dispositif de vérification numérique "
    "(QR code / lien officiel) permettant d’en confirmer l’authenticité et la validité."
)

DECLARATION_TITLE = "Déclaration"
DECLARATION_TEXT = (
    "La présente attestation est délivrée pour servir et valoir ce que de droit auprès de toute autorité "
    "administrative, institutionnelle ou partenaire."
)


def _set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_borders(cell, color: str = "0B4F9C", size: str = "10"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def _add_centered_text(doc: Document, text: str, size: int = 10, bold: bool = False, color: RGBColor | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def _add_article_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = PRIMARY_COLOR
    return p


def _add_normal_paragraph(doc: Document, text: str, justify: bool = True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = TEXT_COLOR
    return p


def _add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(9.3)
    run.font.color.rgb = TEXT_COLOR
    return p


def generate_docx(payload: dict, qr_path: Path) -> Path:
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    ref = payload["reference"]
    out = DOCX_DIR / f"{sanitize_reference(ref)}.docx"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Mm(12)
    section.bottom_margin = Mm(12)
    section.left_margin = Mm(16)
    section.right_margin = Mm(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        p.add_run().add_picture(str(LOGO_PATH), width=Mm(22))

    _add_centered_text(doc, ORG_NAME.upper(), size=14, bold=True, color=PRIMARY_COLOR)
    _add_centered_text(doc, AMALIA_REFERENCE, size=9, bold=False, color=TEXT_COLOR)
    _add_centered_text(doc, ONG_SIEGE, size=9, bold=False, color=TEXT_COLOR)
    _add_centered_text(doc, ONG_REPRESENTATION, size=9, bold=False, color=TEXT_COLOR)

    title_table = doc.add_table(rows=1, cols=1)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    title_cell = title_table.rows[0].cells[0]
    title_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_shading(title_cell, "EAF2FB")
    _set_cell_borders(title_cell, color="0B4F9C", size="12")

    tp = title_cell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("ATTESTATION DE MANDAT OFFICIEL")
    tr.bold = True
    tr.font.size = Pt(14)
    tr.font.color.rgb = PRIMARY_COLOR

    rp = doc.add_paragraph()
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rp.add_run(f"Référence : {ref}")
    rr.bold = True
    rr.font.size = Pt(9.5)
    rr.font.color.rgb = TEXT_COLOR

    if LOGO_PATH.exists():
        wp = doc.add_paragraph()
        wp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        wp.add_run().add_picture(str(LOGO_PATH), width=Mm(32))

    intro_lines = mandate_intro_lines(payload)

    for idx, line in enumerate(intro_lines):
        if not line.strip():
            doc.add_paragraph()
            continue

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if idx not in [7, 8, 13, 15] else WD_ALIGN_PARAGRAPH.LEFT

        run = p.add_run(line)
        run.font.size = Pt(9.5)
        run.font.color.rgb = TEXT_COLOR

        if idx == 6:  # nom du délégué
            run.bold = True
        if idx == 8:  # contact
            run.bold = True
        if idx == 13:  # fonction
            run.bold = True
            run.font.color.rgb = PRIMARY_COLOR

    _add_article_heading(doc, ARTICLE_1_TITLE)
    _add_normal_paragraph(doc, ARTICLE_1_TEXT)

    _add_article_heading(doc, ARTICLE_2_TITLE)
    for item in ARTICLE_2_BULLETS:
        _add_bullet(doc, item)

    _add_article_heading(doc, ARTICLE_3_TITLE)
    _add_normal_paragraph(doc, ARTICLE_3_TEXT)

    _add_article_heading(doc, ARTICLE_4_TITLE)
    _add_normal_paragraph(doc, ARTICLE_4_TEXT)

    _add_article_heading(doc, ARTICLE_5_TITLE)
    article_5 = ARTICLE_5_TEXT_TEMPLATE.format(
        date_emission=payload.get("date_emission", ""),
        date_expiration=payload.get("date_expiration", ""),
    )
    _add_normal_paragraph(doc, article_5)

    _add_article_heading(doc, ARTICLE_6_TITLE)
    _add_normal_paragraph(doc, ARTICLE_6_TEXT)

    _add_article_heading(doc, DECLARATION_TITLE)
    _add_normal_paragraph(doc, DECLARATION_TEXT)

    _add_normal_paragraph(
        doc,
        f"Vérification numérique : {verify_url(payload['reference'], payload['mandate_uid'], payload['signature_token'])}",
        justify=False,
    )

    qp = doc.add_paragraph()
    qp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if qr_path.exists():
        qp.add_run().add_picture(str(qr_path), width=Mm(28))

    sign_table = doc.add_table(rows=2, cols=2)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = sign_table.rows[0].cells
    hdr[0].text = f"Fait à {payload.get('ville_signature', ORG_CITY)}, le {payload.get('date_emission', '')}"
    hdr[1].text = "Pour l’ONG Renaître de Nouveau\nLe Président du Conseil d’Administration"

    row = sign_table.rows[1].cells
    row[0].text = "Cachet officiel"
    row[1].text = ""

    if STAMP_PATH.exists():
        row[0].paragraphs[0].add_run().add_picture(str(STAMP_PATH), width=Mm(22))

    if PRESIDENT_SIGNATURE_PATH.exists():
        row[1].paragraphs[0].add_run().add_picture(str(PRESIDENT_SIGNATURE_PATH), width=Mm(32))
        row[1].add_paragraph(PRESIDENT_NAME)

    doc.save(out)
    return out


def _draw_pdf_header(c: canvas.Canvas, w: float, h: float, ref: str):
    if LOGO_PATH.exists():
        try:
            c.drawImage(
                ImageReader(str(LOGO_PATH)),
                18 * mm,
                h - 30 * mm,
                width=18 * mm,
                height=18 * mm,
                preserveAspectRatio=True,
                mask="auto",
            )
        except Exception:
            pass

    if LOGO_PATH.exists():
        try:
            c.saveState()
            c.setFillAlpha(0.07)
            c.drawImage(
                ImageReader(str(LOGO_PATH)),
                52 * mm,
                85 * mm,
                width=105 * mm,
                height=105 * mm,
                preserveAspectRatio=True,
                mask="auto",
            )
            c.restoreState()
        except Exception:
            pass

    c.setFillColor(PDF_PRIMARY)
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(w / 2, h - 18 * mm, ORG_NAME.upper())

    c.setFillColor(PDF_TEXT)
    c.setFont("Helvetica", 8.7)
    c.drawCentredString(w / 2, h - 24 * mm, AMALIA_REFERENCE)
    c.drawCentredString(w / 2, h - 28.5 * mm, ONG_SIEGE)
    c.drawCentredString(w / 2, h - 33 * mm, ONG_REPRESENTATION)

    box_x = 24 * mm
    box_y = h - 48 * mm
    box_w = w - 48 * mm
    box_h = 11 * mm

    c.setStrokeColor(PDF_PRIMARY)
    c.setFillColor(colors.HexColor("#EAF2FB"))
    c.roundRect(box_x, box_y, box_w, box_h, 3 * mm, stroke=1, fill=1)

    c.setFillColor(PDF_PRIMARY)
    c.setFont("Helvetica-Bold", 13)
    c.drawCentredString(w / 2, box_y + 4 * mm, "ATTESTATION DE MANDAT OFFICIEL")

    c.setFillColor(PDF_TEXT)
    c.setFont("Helvetica-Bold", 9.2)
    c.drawCentredString(w / 2, h - 53 * mm, f"Référence : {ref}")


def _draw_pdf_wrapped(c: canvas.Canvas, text: str, x_mm: float, y: float, width_chars: int, font="Helvetica", size=9.1, leading_mm=4.3):
    c.setFont(font, size)
    c.setFillColor(PDF_TEXT)
    for line in _wrap_text(text, width_chars):
        c.drawString(x_mm * mm, y, line)
        y -= leading_mm * mm
    return y


def _draw_pdf_bullets(c: canvas.Canvas, items: list[str], x_mm: float, y: float, width_chars: int, size=9.0, leading_mm=4.2):
    c.setFont("Helvetica", size)
    c.setFillColor(PDF_TEXT)
    for item in items:
        wrapped = _wrap_text(item, width_chars)
        if not wrapped:
            continue
        c.drawString(x_mm * mm, y, "• " + wrapped[0])
        y -= leading_mm * mm
        for cont in wrapped[1:]:
            c.drawString((x_mm + 4) * mm, y, cont)
            y -= leading_mm * mm
    return y


def generate_pdf(payload: dict, qr_path: Path) -> Path:
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    ref = payload["reference"]
    out = PDF_DIR / f"{sanitize_reference(ref)}.pdf"

    c = canvas.Canvas(str(out), pagesize=A4)
    w, h = A4

    _draw_pdf_header(c, w, h, ref)

    y = h - 61 * mm

    intro_lines = mandate_intro_lines(payload)
    for idx, line in enumerate(intro_lines):
        if not line.strip():
            y -= 2.2 * mm
            continue

        font_name = "Helvetica-Bold" if idx in [6, 8, 13] else "Helvetica"
        font_size = 8.9
        if idx == 13:
            c.setFillColor(PDF_PRIMARY)
        else:
            c.setFillColor(PDF_TEXT)

        for wrapped in _wrap_text(line, 115):
            c.setFont(font_name, font_size)
            c.drawString(18 * mm, y, wrapped)
            y -= 4.0 * mm

    c.setFillColor(PDF_PRIMARY)
    c.setFont("Helvetica-Bold", 9.5)
    c.drawString(18 * mm, y, ARTICLE_1_TITLE)
    y -= 4.2 * mm
    y = _draw_pdf_wrapped(c, ARTICLE_1_TEXT, 18, y, 118, size=8.8, leading_mm=4.0)
    y -= 1.2 * mm

    c.setFillColor(PDF_PRIMARY)
    c.setFont("Helvetica-Bold", 9.5)
    c.drawString(18 * mm, y, ARTICLE_2_TITLE)
    y -= 4.2 * mm
    y = _draw_pdf_bullets(c, ARTICLE_2_BULLETS, 18, y, 112, size=8.8, leading_mm=3.9)
    y -= 1.0 * mm

    c.setFillColor(PDF_PRIMARY)
    c.setFont("Helvetica-Bold", 9.5)
    c.drawString(18 * mm, y, ARTICLE_3_TITLE)
    y -= 4.2 * mm
    y = _draw_pdf_wrapped(c, ARTICLE_3_TEXT, 18, y, 118, size=8.8, leading_mm=4.0)
    y -= 1.0 * mm

    c.setFillColor(PDF_PRIMARY)
    c.setFont("Helvetica-Bold", 9.5)
    c.drawString(18 * mm, y, ARTICLE_4_TITLE)
    y -= 4.2 * mm
    y = _draw_pdf_wrapped(c, ARTICLE_4_TEXT, 18, y, 118, size=8.8, leading_mm=4.0)
    y -= 1.0 * mm

    c.setFillColor(PDF_PRIMARY)
    c.setFont("Helvetica-Bold", 9.5)
    c.drawString(18 * mm, y, ARTICLE_5_TITLE)
    y -= 4.2 * mm
    article_5 = ARTICLE_5_TEXT_TEMPLATE.format(
        date_emission=payload.get("date_emission", ""),
        date_expiration=payload.get("date_expiration", ""),
    )
    y = _draw_pdf_wrapped(c, article_5, 18, y, 118, size=8.8, leading_mm=4.0)
    y -= 1.0 * mm

    c.setFillColor(PDF_PRIMARY)
    c.setFont("Helvetica-Bold", 9.5)
    c.drawString(18 * mm, y, ARTICLE_6_TITLE)
    y -= 4.2 * mm
    y = _draw_pdf_wrapped(c, ARTICLE_6_TEXT, 18, y, 118, size=8.8, leading_mm=4.0)
    y -= 1.0 * mm

    c.setFillColor(PDF_PRIMARY)
    c.setFont("Helvetica-Bold", 9.5)
    c.drawString(18 * mm, y, DECLARATION_TITLE)
    y -= 4.2 * mm
    y = _draw_pdf_wrapped(c, DECLARATION_TEXT, 18, y, 118, size=8.8, leading_mm=4.0)

    if qr_path.exists():
        try:
            c.drawImage(
                ImageReader(str(qr_path)),
                18 * mm,
                16 * mm,
                width=24 * mm,
                height=24 * mm,
                preserveAspectRatio=True,
                mask="auto",
            )
        except Exception:
            pass

    c.setFont("Helvetica", 7.1)
    c.setFillColor(PDF_TEXT)
    verif = verify_url(payload["reference"], payload["mandate_uid"], payload["signature_token"])
    y_ver = 34 * mm
    x_ver = 46 * mm
    max_lines = 3
    wrapped_verif = _wrap_text("Vérification numérique : " + verif, 62)
    wrapped_verif = wrapped_verif[:max_lines]
    for i, line in enumerate(wrapped_verif):
        c.drawString(x_ver, y_ver - i * 3.6 * mm, line)

    c.setFont("Helvetica", 8.5)
    c.drawString(118 * mm, 28 * mm, f"Fait à {payload.get('ville_signature', ORG_CITY)}, le {payload.get('date_emission', '')}")

    if STAMP_PATH.exists():
        try:
            c.drawImage(
                ImageReader(str(STAMP_PATH)),
                118 * mm,
                10 * mm,
                width=20 * mm,
                height=20 * mm,
                preserveAspectRatio=True,
                mask="auto",
            )
        except Exception:
            pass

    if PRESIDENT_SIGNATURE_PATH.exists():
        try:
            c.drawImage(
                ImageReader(str(PRESIDENT_SIGNATURE_PATH)),
                147 * mm,
                14 * mm,
                width=32 * mm,
                height=12 * mm,
                preserveAspectRatio=True,
                mask="auto",
            )
        except Exception:
            pass

    c.setFont("Helvetica-Bold", 8.5)
    c.drawString(145 * mm, 12 * mm, PRESIDENT_NAME)
    c.setFont("Helvetica", 7.8)
    c.drawString(145 * mm, 8 * mm, PRESIDENT_TITLE)

    c.save()
    return out


def _wrap_text(text: str, width: int) -> list[str]:
    words = text.split()
    lines = []
    current = []
    count = 0

    for word in words:
        extra = 1 if current else 0
        if count + len(word) + extra <= width:
            current.append(word)
            count += len(word) + extra
        else:
            lines.append(" ".join(current))
            current = [word]
            count = len(word)

    if current:
        lines.append(" ".join(current))

    return lines


def generate_all(payload: dict):
    payload["signature_token"] = generate_signature(
        payload["reference"],
        payload["mandate_uid"],
    )

    qr = generate_qr(payload["reference"], payload["mandate_uid"], payload["signature_token"])
    docx = generate_docx(payload, qr)
    pdf = generate_pdf(payload, qr)
    return qr, docx, pdf
